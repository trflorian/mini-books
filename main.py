from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

flip_map = {
    'a':'ɐ','b':'q','c':'ɔ','d':'p','e':'ǝ','f':'ɟ','g':'ƃ','h':'ɥ','i':'ı',
    'j':'ɾ','k':'ʞ','l':'ʃ','m':'ɯ','n':'u','o':'o','p':'d','q':'b','r':'ɹ',
    's':'s','t':'ʇ','u':'n','v':'ʌ','w':'ʍ','x':'x','y':'ʎ','z':'z',
    'A':'∀','B':'𐐒','C':'Ɔ','D':'◖','E':'Ǝ','F':'Ⅎ','G':'⅁','H':'H','I':'I',
    'J':'ſ','K':'⋊','L':'Ꞁ','M':'W','N':'N','O':'O','P':'Ԁ','Q':'Q','R':'ɹ',
    'S':'S','T':'⊥','U':'∩','V':'Λ','W':'M','X':'X','Y':'⅄','Z':'Z',
    '0':'0','1':'Ɩ','2':'ᄅ','3':'Ɛ','4':'ㄣ','5':'ϛ','6':'9','7':'ㄥ','8':'8','9':'6',
    ',':"'", '.':',','?':'¿','!':'¡','"':'„',"'":',','(' : ')', ')':'(','[':']',']':'[','{':'}','}':'{','<':'>','>':'<','&':'⅋','_':'‾'
}

def flip_text(s: str) -> str:
    # reverse string, then map each char
    return ''.join(flip_map.get(c, c) for c in s[::-1])

document = Document()

section = document.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Mm(297)
section.page_height = Mm(210)
section.left_margin = Mm(4.0)
section.right_margin = Mm(4.0)
section.top_margin = Mm(4.0)
section.bottom_margin = Mm(0.0)

table = document.add_table(rows=2, cols=4)

cell_indices = [[1, 8, 7, 6], [2, 3, 4, 5]]
cell_flips = [4 * [True], 4 * [False]]

for row_id, row in enumerate(table.rows):
    row.height = Mm(100)
    for col_id, cell in enumerate(row.cells):
        cell_id = cell_indices[row_id][col_id]
        cell_flip = cell_flips[row_id][col_id]

        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        cell_text = f"This is a text cell in row {row_id + 1}, column {col_id + 1}. The mini-book page number is {cell_id}."

        if cell_flip:
            cell.text = flip_text(cell_text)
        else:
            cell.text = cell_text

        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

document.add_page_break()

document.save("demo.docx")
