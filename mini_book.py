import re
import tempfile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm
from docx2pdf import convert

from utils import CellDirection, create_image_for_direction, set_vertical_cell_direction


class MiniBook:
    """
    A class to create a mini book layout in a Word document.
    """

    # A4 Portrait Page Setup, all values in millimeters
    _PAGE_ORIENTATION = WD_ORIENT.PORTRAIT
    _PAGE_WIDTH = 210
    _PAGE_HEIGHT = 297
    _MARGIN = 4.0
    _PARAGRAPH_MARGIN = 8.0

    _NUM_PAGES = 8
    _DEFAULT_PAGE_CONTENT = [str(i) for i in range(_NUM_PAGES)]

    def __init__(self, page_content: list[str] = _DEFAULT_PAGE_CONTENT) -> None:
        if len(page_content) != self._NUM_PAGES:
            raise ValueError(
                f"Page content must contain exactly {self._NUM_PAGES} items."
            )

        self._document = Document()

        self._setup_page()
        self._setup_table(page_content)

    def _setup_page(self) -> None:
        page_section = self._document.sections[0]

        # Page Setup
        page_section.orientation = self._PAGE_ORIENTATION
        page_section.page_width = Mm(self._PAGE_WIDTH)
        page_section.page_height = Mm(self._PAGE_HEIGHT)

        # Margins
        page_section.left_margin = Mm(self._MARGIN)
        page_section.right_margin = Mm(self._MARGIN)
        page_section.top_margin = Mm(self._MARGIN)
        page_section.bottom_margin = Mm(self._MARGIN)

        # paragraph format
        paragraph_format = self._document.styles['Normal'].paragraph_format
        paragraph_format.left_indent = Mm(self._PARAGRAPH_MARGIN)
        paragraph_format.right_indent = Mm(self._PARAGRAPH_MARGIN)

        # Font
        font = self._document.styles['Normal'].font
        font.name = "Arial"
        font.size = Mm(6.0)

    def _setup_table(self, page_content: list[str]) -> None:
        # Table Setup
        table = self._document.add_table(rows=4, cols=2)

        cell_indices = [[2, 1], [3, 8], [4, 7], [5, 6]]
        cell_orientations = 4 * [[CellDirection.TB_RL, CellDirection.BT_LR]]

        for row_id, row in enumerate(table.rows):
            row.height = Mm(self._PAGE_HEIGHT / len(table.rows) - self._MARGIN)
            for col_id, cell in enumerate(row.cells):
                cell_id = cell_indices[row_id][col_id]
                cell_orientation = cell_orientations[row_id][col_id]

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                paragraph = cell.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run()

                content = page_content[cell_id - 1]
                img_match = re.match("<img>(.*?)</img>", content)
                if img_match:
                    # Extract image path from content
                    img_path = img_match.group(1).strip()
                    with create_image_for_direction(
                        image_path=img_path,
                        direction=cell_orientation,
                    ) as rotated_img_path:
                        run.add_picture(rotated_img_path, width=Mm(55.0), height=Mm(55.0))
                else:
                    run.add_text(content)
                    
                    set_vertical_cell_direction(cell, cell_orientation)

    def save(self, filename: str) -> None:
        self._document.save(filename)

    def export_to_pdf(self, filename: str) -> None:
        """
        Export the mini book to a PDF file.
        NOTE: Requires `docx2pdf` package, only works on Windows and macOS.
        Args:
            filename (str): The name of the output PDF file.
        """
        with tempfile.NamedTemporaryFile(suffix=".docx") as tmp_file:
            tmp_filename = tmp_file.name
            self._document.save(tmp_filename)

            convert(
                input_path=tmp_filename,
                output_path=filename,
            )
